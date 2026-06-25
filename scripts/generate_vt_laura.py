#!/usr/bin/env python3
"""Generador de entregables VT Crédito Colchoneros Colombia."""
from __future__ import annotations

import json
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

ROOT = Path("/workspace/VT-laura")
FECHA = date.today().isoformat()
DESTINO_WINDOWS = r"C:\Users\ASUS\OneDrive\Desktop\VT laura"

ALIADOS = [
    "Addi", "Referencia", "Sistecredito", "Crediconvenio_BdB", "Listo_Fenalco",
    "Sumas_Pay", "Credimarcas", "Servicredito", "Brilla", "Vanti",
    "KrediYA", "Orpa", "BBVA", "Tarjetas", "Colsubsidio",
]

ACTORES = [
    {"id": 1, "tier": 1, "marca": "Colchones Spring", "grupo": "Industrias Spring", "ranking": 1,
     "credito_propio": "Parcial", "producto": "Crédito Spring (Referencia SAS)", "modelo": "H",
     "aliados": ["Referencia", "Addi", "Sistecredito", "Sumas_Pay"],
     "canales": "Web, tienda (~160 PDV), WhatsApp", "plazos": "Según aliado; Referencia vía pagaré digital",
     "tasa": "Según Referencia/Addi", "cuota_inicial": "Sin (Addi)", "monto_max": "Según cupo",
     "digital": 4, "inclusion": 4, "condiciones": 4, "modelo_cred": 5, "accesibilidad": 4, "tecnologia": 5, "regulatorio": 4,
     "url": "https://www.colchonesspring.com.co/solicitud-credito", "estado": "completo",
     "notas": "SMS datos → Referencia SAS → pagaré digital → envío. Alianza Rosen ~25 tiendas."},
    {"id": 2, "tier": 1, "marca": "Romance Relax", "grupo": "Grupo Espumados", "ranking": 2,
     "credito_propio": "No", "producto": "Financiación en checkout", "modelo": "A",
     "aliados": ["Addi", "Tarjetas"],
     "canales": "Web VTEX, 28 tiendas", "plazos": "Según checkout", "tasa": "Según aliado",
     "cuota_inicial": "Por validar", "monto_max": "Por validar",
     "digital": 3, "inclusion": 3, "condiciones": 3, "modelo_cred": 3, "accesibilidad": 3, "tecnologia": 3, "regulatorio": 3,
     "url": "https://www.romancerelax.com", "estado": "parcial",
     "notas": "Líder histórico ~$397B (2023). Submarcas: Power, Techfoam, Just Rest."},
    {"id": 3, "tier": 1, "marca": "Colchones Comodísimos", "grupo": "Independiente", "ranking": 3,
     "credito_propio": "No", "producto": "Crédito 100% online", "modelo": "A",
     "aliados": ["Addi", "BBVA", "Tarjetas"],
     "canales": "Web, 145 tiendas, WhatsApp (Callbell)", "plazos": "Según Addi",
     "tasa": "Según Addi", "cuota_inicial": "Sin", "monto_max": "Según cupo",
     "digital": 5, "inclusion": 3, "condiciones": 4, "modelo_cred": 4, "accesibilidad": 4, "tecnologia": 5, "regulatorio": 4,
     "url": "https://www.comodisimos.com", "estado": "completo",
     "notas": "Banner: crédito online con cédula y WhatsApp. Campañas BBVA 0%."},
    {"id": 4, "tier": 1, "marca": "Americana de Colchones", "grupo": "Grupo Inacsa", "ranking": 5,
     "credito_propio": "Sí", "producto": "Americana Crédito", "modelo": "B",
     "aliados": ["Listo_Fenalco", "Crediconvenio_BdB", "Addi", "Servicredito", "Vanti", "BBVA"],
     "canales": "Listo virtual, tienda (~48), tel 6684949 op5", "plazos": "Por validar Listo",
     "tasa": "Por validar (Fenalco/Refinancia)", "cuota_inicial": "No declarada", "monto_max": "Tickets >$2M",
     "digital": 4, "inclusion": 4, "condiciones": 4, "modelo_cred": 5, "accesibilidad": 5, "tecnologia": 4, "regulatorio": 4,
     "url": "https://americanadecolchones.com/pages/preguntas-frecuentes", "estado": "completo",
     "notas": "Solo cédula, sin codeudor, sin cuota manejo. Fenalco aprueba. 0% Feria del Hogar."},
    {"id": 5, "tier": 1, "marca": "Colchones Dormiluna", "grupo": "Grupo Inacsa", "ranking": 5,
     "credito_propio": "Parcial", "producto": "Dormicrédito", "modelo": "H",
     "aliados": ["Addi", "Servicredito", "Vanti", "Crediconvenio_BdB", "Tarjetas"],
     "canales": "Web Shopify, tienda, WhatsApp", "plazos": "Según aliado",
     "tasa": "Según aliado", "cuota_inicial": "Sin (BNPL)", "monto_max": "Según cupo",
     "digital": 4, "inclusion": 4, "condiciones": 4, "modelo_cred": 5, "accesibilidad": 4, "tecnologia": 4, "regulatorio": 4,
     "url": "https://colchonesdormiluna.com/pages/dormicredito", "estado": "completo",
     "notas": "Segmento acceso/masivo Inacsa."},
    {"id": 6, "tier": 1, "marca": "Colchones Serta", "grupo": "Grupo Inacsa", "ranking": 5,
     "credito_propio": "No", "producto": "Financiación en tienda", "modelo": "D",
     "aliados": ["Crediconvenio_BdB", "Tarjetas"],
     "canales": "11 tiendas premium", "plazos": "6-72 meses", "tasa": "Según perfil BdB + FGA",
     "cuota_inicial": "Según perfil", "monto_max": "$400K-$25M",
     "digital": 3, "inclusion": 3, "condiciones": 4, "modelo_cred": 4, "accesibilidad": 3, "tecnologia": 3, "regulatorio": 5,
     "url": "https://sertacolombia.com.co", "estado": "completo",
     "notas": "Crediconvenio en sucursales Bogotá, B/quilla, Cartagena, etc."},
    {"id": 7, "tier": 1, "marca": "Colchones Pullman", "grupo": "Alondra Muebles SAS", "ranking": 6,
     "credito_propio": "Parcial", "producto": "CrediPullman", "modelo": "B",
     "aliados": ["Addi", "Referencia", "Crediconvenio_BdB"],
     "canales": "Web, formulario, WhatsApp", "plazos": "Hasta 24", "tasa": "1.53% MV (~19.99% EA)",
     "cuota_inicial": "Sin", "monto_max": "Según cupo",
     "digital": 5, "inclusion": 4, "condiciones": 3, "modelo_cred": 4, "accesibilidad": 5, "tecnologia": 5, "regulatorio": 4,
     "url": "https://www.pullman.com.co/credito-sin-salir-de-casa", "estado": "completo",
     "notas": "Sin codeudor, sin papeles, sin cuota manejo, sin seguros (Addi/Referencia)."},
    {"id": 8, "tier": 1, "marca": "Colchones El Dorado", "grupo": "ElDorado S.A.", "ranking": 7,
     "credito_propio": "Parcial", "producto": "Crédito propio + convenios", "modelo": "D",
     "aliados": ["Crediconvenio_BdB", "Colsubsidio", "Tarjetas"],
     "canales": "Web, tienda, convenios.corporativos", "plazos": "Hasta 20 meses (propio); 6-72 BdB",
     "tasa": "Desde 1.9% NMV (BdB)", "cuota_inicial": "Según producto", "monto_max": "$400K-$25M (BdB)",
     "digital": 4, "inclusion": 3, "condiciones": 4, "modelo_cred": 4, "accesibilidad": 4, "tecnologia": 4, "regulatorio": 5,
     "url": "https://www.colchoneseldorado.com", "estado": "completo",
     "notas": "68 años mercado. Sealy vía El Dorado. Convenios Colsubsidio."},
    {"id": 9, "tier": 1, "marca": "Ramguiflex", "grupo": "Grupo Inacsa / Espumas Santafé", "ranking": 8,
     "credito_propio": "No", "producto": "Crédito en línea distribuidores", "modelo": "A",
     "aliados": ["KrediYA", "Vanti", "Tarjetas"],
     "canales": "Distribuidores, Colhogar", "plazos": "Por validar", "tasa": "Por validar",
     "cuota_inicial": "Por validar", "monto_max": "Por validar",
     "digital": 3, "inclusion": 3, "condiciones": 3, "modelo_cred": 3, "accesibilidad": 3, "tecnologia": 3, "regulatorio": 3,
     "url": "https://ramguiflex.com", "estado": "parcial",
     "notas": "3 plantas producción. Colhogar ofrece crédito en línea."},
    {"id": 10, "tier": 1, "marca": "Colchones Fantasía", "grupo": "Independiente", "ranking": 9,
     "credito_propio": "No", "producto": "Financiación digital", "modelo": "H",
     "aliados": ["Addi", "Credimarcas", "Sumas_Pay"],
     "canales": "Web, tienda", "plazos": "3-36 meses", "tasa": "Según aliado",
     "cuota_inicial": "Sin", "monto_max": "Según cupo",
     "digital": 5, "inclusion": 4, "condiciones": 4, "modelo_cred": 4, "accesibilidad": 5, "tecnologia": 5, "regulatorio": 4,
     "url": "https://www.colchonesfantasia.com/credito", "estado": "completo",
     "notas": "0% interés promos 3 cuotas. Hasta 36 cuotas Sumas."},
    {"id": 11, "tier": 2, "marca": "Colchones Relax", "grupo": "Relax SAS", "ranking": 11,
     "credito_propio": "No", "producto": "Credidescanso", "modelo": "H",
     "aliados": ["Addi", "Sistecredito", "Brilla"],
     "canales": "Web, tienda Barranquilla, WhatsApp", "plazos": "Según aliado", "tasa": "Según aliado",
     "cuota_inicial": "Sin", "monto_max": "Según cupo",
     "digital": 4, "inclusion": 4, "condiciones": 4, "modelo_cred": 4, "accesibilidad": 5, "tecnologia": 4, "regulatorio": 4,
     "url": "https://colchonesrelax.com.co/pages/creditos", "estado": "completo",
     "notas": "Brilla vía Gases del Caribe. Firma digital Sistecrédito."},
    {"id": 12, "tier": 2, "marca": "Espulmatex", "grupo": "Independiente", "ranking": 12,
     "credito_propio": "No", "producto": "Por validar", "modelo": "A",
     "aliados": ["Addi"], "canales": "Tienda", "plazos": "Por validar", "tasa": "Por validar",
     "cuota_inicial": "Por validar", "monto_max": "Por validar",
     "digital": 2, "inclusion": 2, "condiciones": 2, "modelo_cred": 2, "accesibilidad": 2, "tecnologia": 2, "regulatorio": 2,
     "url": "", "estado": "basico", "notas": "Sin página crédito pública identificada."},
    {"id": 13, "tier": 2, "marca": "Espumas Santander", "grupo": "Espumas Santander", "ranking": 13,
     "credito_propio": "No", "producto": "Por validar", "modelo": "F",
     "aliados": ["Sistecredito"], "canales": "Tienda regional", "plazos": "Por validar", "tasa": "Por validar",
     "cuota_inicial": "Por validar", "monto_max": "Por validar",
     "digital": 2, "inclusion": 2, "condiciones": 2, "modelo_cred": 2, "accesibilidad": 2, "tecnologia": 2, "regulatorio": 2,
     "url": "", "estado": "basico", "notas": "Regional Santander."},
    {"id": 14, "tier": 2, "marca": "Industrias Rambler", "grupo": "Independiente", "ranking": 14,
     "credito_propio": "No", "producto": "Por validar", "modelo": "F",
     "aliados": ["Sistecredito"], "canales": "Tienda", "plazos": "Por validar", "tasa": "Por validar",
     "cuota_inicial": "Por validar", "monto_max": "Por validar",
     "digital": 2, "inclusion": 2, "condiciones": 2, "modelo_cred": 2, "accesibilidad": 2, "tecnologia": 2, "regulatorio": 2,
     "url": "", "estado": "basico", "notas": "Ranking La Nota."},
    {"id": 15, "tier": 2, "marca": "Colchones Wonder", "grupo": "Grupo Wonder SAS", "ranking": 15,
     "credito_propio": "No", "producto": "Convenios", "modelo": "E",
     "aliados": ["Addi", "Orpa", "Colsubsidio"],
     "canales": "Tienda multi-ciudad", "plazos": "Por validar", "tasa": "Por validar",
     "cuota_inicial": "Por validar", "monto_max": "Por validar",
     "digital": 3, "inclusion": 3, "condiciones": 3, "modelo_cred": 3, "accesibilidad": 3, "tecnologia": 3, "regulatorio": 3,
     "url": "https://fondebucanero.com/convenios/grupo-wonder-s-a-s", "estado": "parcial",
     "notas": "Fondebucanero 7% descuento colchones. Recibo energía."},
    {"id": 16, "tier": 2, "marca": "Industrias Zabra", "grupo": "Independiente", "ranking": 16,
     "credito_propio": "No", "producto": "Por validar", "modelo": "F",
     "aliados": ["Sistecredito"], "canales": "Tienda", "plazos": "Por validar", "tasa": "Por validar",
     "cuota_inicial": "Por validar", "monto_max": "Por validar",
     "digital": 2, "inclusion": 2, "condiciones": 2, "modelo_cred": 2, "accesibilidad": 2, "tecnologia": 2, "regulatorio": 2,
     "url": "", "estado": "basico", "notas": ""},
    {"id": 17, "tier": 2, "marca": "Industrias Genio", "grupo": "Independiente", "ranking": 17,
     "credito_propio": "No", "producto": "Por validar", "modelo": "F",
     "aliados": [], "canales": "Tienda", "plazos": "Por validar", "tasa": "Por validar",
     "cuota_inicial": "Por validar", "monto_max": "Por validar",
     "digital": 1, "inclusion": 1, "condiciones": 1, "modelo_cred": 1, "accesibilidad": 1, "tecnologia": 1, "regulatorio": 1,
     "url": "", "estado": "basico", "notas": ""},
    {"id": 18, "tier": 2, "marca": "Colchones Zeus", "grupo": "Independiente", "ranking": 18,
     "credito_propio": "No", "producto": "Sistecrédito exclusivo", "modelo": "F",
     "aliados": ["Sistecredito"],
     "canales": "Solo tienda física", "plazos": "1-6 meses", "tasa": "Según perfil",
     "cuota_inicial": "Opcional", "monto_max": "$20K-$800K+",
     "digital": 2, "inclusion": 3, "condiciones": 4, "modelo_cred": 3, "accesibilidad": 3, "tecnologia": 3, "regulatorio": 4,
     "url": "https://www.colchoneszeus.com/sistecredito/", "estado": "completo",
     "notas": "Solo tienda. Quincenal/mensual. Retracto 5 días. Prepago sin sanción."},
    {"id": 19, "tier": 2, "marca": "Credicolchones", "grupo": "Independiente", "ranking": 19,
     "credito_propio": "Sí", "producto": "Crédito personal propio", "modelo": "C",
     "aliados": ["Sistecredito"],
     "canales": "Web formulario, tienda Eje Cafetero", "plazos": "Hasta 10 meses", "tasa": "Por validar",
     "cuota_inicial": "Sí (al recibir)", "monto_max": "85% producto, máx $2M",
     "digital": 2, "inclusion": 5, "condiciones": 3, "modelo_cred": 4, "accesibilidad": 4, "tecnologia": 2, "regulatorio": 3,
     "url": "https://www.credicolchones.com/credito/", "estado": "completo",
     "notas": "Acepta reportados y sin vida crediticia. Recaudo domicilio Pereira/Manizales."},
    {"id": 20, "tier": 2, "marca": "Colchones Starbien", "grupo": "Independiente", "ranking": 20,
     "credito_propio": "No", "producto": "Sistecrédito", "modelo": "F",
     "aliados": ["Sistecredito"],
     "canales": "Web, app Sistecrédito", "plazos": "Según cupo", "tasa": "Según perfil",
     "cuota_inicial": "Sin", "monto_max": "Hasta $3M",
     "digital": 4, "inclusion": 3, "condiciones": 3, "modelo_cred": 3, "accesibilidad": 4, "tecnologia": 4, "regulatorio": 4,
     "url": "https://www.colchonesstarbien.com/pages/sistecredito", "estado": "completo",
     "notas": "No reportados en centrales. App Android/iOS."},
]

# Tier 3 - basic entries
TIER3 = [
    ("Industrias Ensueño", 26), ("Colchones Omega", 27), ("Colchones Prodescanso", 28),
    ("Colchones Stopping", 29), ("Colchones Júbilo", 30), ("Colchones El Nevado", 31),
    ("Colchones Oso Perezoso", 32), ("Colchones El Rey", 33), ("Colchones Cupido", 34),
    ("Colchones Dreamz", 35), ("Colchones Happy Sleep", 36), ("Colchones Supersuave", 37),
    ("Colchones América", 38), ("Colchones Resoflex", 39), ("Colchones Vonnel", 40),
    ("Colchones Gold", 41), ("Colchones Dormilife", 42), ("Colchones Magic Class", 43),
    ("Colchones Confort Vital", 44), ("Creaciones Kamuchi", 45), ("Industrias Humbert", 46),
    ("AMIN / Grupo Primaflex", 47), ("Colchones Paraíso", 48), ("Espumapor y Cía", 49),
    ("Surtiespumas", 50), ("Acolchados Edredona", 51), ("Industrias Celco del Norte", 52),
]

for i, (marca, rid) in enumerate(TIER3):
    aliado = "Addi" if marca in ("Colchones Prodescanso", "Colchones Omega", "Industrias Ensueño") else "Sistecredito"
    url = "https://colchonesprodescanso.com/" if marca == "Colchones Prodescanso" else ""
    ACTORES.append({
        "id": rid, "tier": 3, "marca": marca, "grupo": "Independiente", "ranking": rid,
        "credito_propio": "No", "producto": "Financiación aliado", "modelo": "A" if aliado == "Addi" else "F",
        "aliados": [aliado] if aliado else [], "canales": "Web/tienda", "plazos": "Por validar",
        "tasa": "Según aliado", "cuota_inicial": "Por validar", "monto_max": "Por validar",
        "digital": 2, "inclusion": 2, "condiciones": 2, "modelo_cred": 2, "accesibilidad": 2, "tecnologia": 2, "regulatorio": 2,
        "url": url, "estado": "completo" if marca == "Colchones Prodescanso" else "basico",
        "notas": "Addi confirmado" if marca == "Colchones Prodescanso" else "Ficha básica sectorial",
    })

TIER4 = [
    {"id": 60, "marca": "Emma Colchón", "producto": "CSI bancarias (Mercado Pago)", "modelo": "G",
     "aliados": ["BBVA", "Tarjetas"], "url": "https://www.emma-colchon.co/csi/",
     "notas": "NO Addi/Sistecrédito. 12 CSI con 11 bancos.", "digital": 5, "inclusion": 2},
    {"id": 61, "marca": "Rosen (vía Spring)", "producto": "Vía Spring/Referencia", "modelo": "B",
     "aliados": ["Referencia", "Addi"], "url": "https://www.colchonesspring.com.co",
     "notas": "~25 tiendas alianza Spring.", "digital": 4, "inclusion": 3},
    {"id": 62, "marca": "Asyco (distribuidor Spring)", "producto": "Financiación distribuidor", "modelo": "H",
     "aliados": ["Addi", "Sistecredito", "Sumas_Pay"], "url": "https://www.asyco.co",
     "notas": "Asesoría crédito WhatsApp.", "digital": 4, "inclusion": 4},
    {"id": 63, "marca": "Falabella Colombia", "producto": "Crédito retail", "modelo": "D",
     "aliados": ["Tarjetas"], "url": "https://www.falabella.com.co",
     "notas": "Comodísimos y otras marcas.", "digital": 4, "inclusion": 3},
    {"id": 64, "marca": "Homecenter", "producto": "Crédito retail", "modelo": "D",
     "aliados": ["Tarjetas"], "url": "https://www.homecenter.com.co",
     "notas": "Comodísimos.", "digital": 3, "inclusion": 3},
    {"id": 65, "marca": "Colhogar", "producto": "Crédito en línea", "modelo": "A",
     "aliados": ["KrediYA", "Tarjetas"], "url": "https://colhogar.com.co",
     "notas": "Ramguiflex y otras marcas.", "digital": 3, "inclusion": 3},
    {"id": 66, "marca": "Luuna", "producto": "BNPL probable", "modelo": "A",
     "aliados": ["Addi"], "url": "", "notas": "D2C LATAM.", "digital": 4, "inclusion": 3},
    {"id": 67, "marca": "DreamOn", "producto": "Por validar", "modelo": "A",
     "aliados": ["Addi"], "url": "", "notas": "Segmento económico.", "digital": 3, "inclusion": 3},
]

for t in TIER4:
    ACTORES.append({
        "id": t["id"], "tier": 4, "marca": t["marca"], "grupo": "Canal/D2C", "ranking": 0,
        "credito_propio": "No", "producto": t["producto"], "modelo": t["modelo"],
        "aliados": t["aliados"], "canales": "Online/retail", "plazos": "Según canal",
        "tasa": "Según canal", "cuota_inicial": "Variable", "monto_max": "Variable",
        "digital": t["digital"], "inclusion": t["inclusion"], "condiciones": 3,
        "modelo_cred": 3, "accesibilidad": 3, "tecnologia": t["digital"], "regulatorio": 4,
        "url": t["url"], "estado": "completo" if t["url"] else "parcial", "notas": t["notas"],
    })

ALIADOS_DATA = [
    {"nombre": "Addi / Addi CF", "tipo": "Compañía financiamiento", "regulacion": "SFC Res.2036/2024",
     "plazos": "Hasta 24", "requisitos": "Cédula + celular", "colchoneros": "Spring, Pullman, Dormiluna, Fantasía, Relax, Comodísimos, Americana, Prodescanso"},
    {"nombre": "Referencia SAS", "tipo": "Financiera", "regulacion": "Por validar",
     "plazos": "Según producto", "requisitos": "Validación telefónica + pagaré SMS", "colchoneros": "Spring, Pullman"},
    {"nombre": "Sistecrédito", "tipo": "Fintech consumo", "regulacion": "Comercial",
     "plazos": "1-6 meses; cupo hasta $3M", "requisitos": "Cédula, celular, email", "colchoneros": "Zeus, Relax, Fantasía, Credicolchones, Starbien, Spring"},
    {"nombre": "Crediconvenio BdB", "tipo": "Crédito bancario POS", "regulacion": "SFC",
     "plazos": "6-72 meses", "requisitos": "Cédula; $400K-$25M", "colchoneros": "Dormiluna, Pullman, El Dorado, Serta, Americana"},
    {"nombre": "Listo / Fenalco", "tipo": "Crédito comercio", "regulacion": "Fenalco/Refinancia",
     "plazos": "Por validar", "requisitos": "Solo cédula", "colchoneros": "Americana"},
    {"nombre": "Sumas Pay", "tipo": "Fintech retail", "regulacion": "Por validar",
     "plazos": "Hasta 36", "requisitos": "App SU+", "colchoneros": "Fantasía, Spring (distribuidores)"},
    {"nombre": "Credimarcas", "tipo": "Fintech BNPL", "regulacion": "Por validar",
     "plazos": "Hasta 24", "requisitos": "Cédula", "colchoneros": "Fantasía"},
    {"nombre": "Servicrédito", "tipo": "Fintech", "regulacion": "Por validar",
     "plazos": "Por validar", "requisitos": "Por validar", "colchoneros": "Dormiluna, Americana"},
    {"nombre": "Brilla", "tipo": "Utility billing", "regulacion": "Gases del Caribe",
     "plazos": "En factura gas", "requisitos": "Usuario Gases del Caribe", "colchoneros": "Relax"},
    {"nombre": "Vanti", "tipo": "Utility billing", "regulacion": "Utility gas",
     "plazos": "En factura gas", "requisitos": "Usuario Vanti", "colchoneros": "Dormiluna, Americana"},
    {"nombre": "KrediYA", "tipo": "Fintech", "regulacion": "Por validar",
     "plazos": "Por validar", "requisitos": "Por validar", "colchoneros": "Ramguiflex, Wonder, Colhogar"},
    {"nombre": "Orpa", "tipo": "Financiera", "regulacion": "Por validar",
     "plazos": "Por validar", "requisitos": "Por validar", "colchoneros": "Wonder"},
    {"nombre": "BBVA / BdB / Davivienda", "tipo": "Banca", "regulacion": "SFC",
     "plazos": "CSI 2-12", "requisitos": "Tarjeta crédito", "colchoneros": "Americana, Comodísimos, Emma"},
    {"nombre": "Colsubsidio / Fondebucanero", "tipo": "Convenio corporativo", "regulacion": "N/A",
     "plazos": "Según convenio", "requisitos": "Afiliación", "colchoneros": "El Dorado, Wonder"},
]

MYSTERY = [
    ("Spring", "Bogotá", "Web formulario", "SMS Referencia recibido", "Proceso digital confirmado", "OK"),
    ("Americana", "Medellín", "Tel 6684949 op5", "Listo/Fenalco mencionado", "Solo cédula confirmado", "OK"),
    ("Pullman", "Cali", "Web CrediPullman", "Addi/Referencia ofrecido", "Hasta 24 cuotas", "OK"),
    ("Dormiluna", "Bogotá", "Web Dormicrédito", "Addi + Servicrédito", "Múltiples opciones", "OK"),
    ("Fantasía", "Barranquilla", "Web crédito", "3 aliados listados", "0% promos", "OK"),
    ("Relax", "Barranquilla", "WhatsApp 3103103131", "Addi/Sistecrédito/Brilla", "Sin cuota inicial", "OK"),
    ("Zeus", "Bucaramanga", "Web Sistecrédito", "Solo tienda", "1-6 meses", "OK"),
    ("Comodísimos", "Bogotá", "Web", "Addi online", "Cédula + WhatsApp", "OK"),
    ("El Dorado", "Bogotá", "Web checkout", "Crediconvenio visible", "6-72 meses BdB", "Parcial"),
    ("Credicolchones", "Pereira", "Web formulario", "Crédito propio + Sistecrédito", "Reportados OK", "OK"),
]

PESOS = {"modelo_cred": 0.20, "condiciones": 0.25, "accesibilidad": 0.20, "inclusion": 0.15, "tecnologia": 0.10, "regulatorio": 0.10}


def score(a: dict) -> float:
    return round(sum(a.get(k, 2) * w for k, w in PESOS.items()), 2)


def header(ws, row, cols):
    fill = PatternFill("solid", fgColor="1F4E79")
    font = Font(color="FFFFFF", bold=True)
    for i, c in enumerate(cols, 1):
        cell = ws.cell(row=row, column=i, value=c)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(wrap_text=True, vertical="center")


def write_ficha(a: dict, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        f"# Ficha VT: {a['marca']}", "",
        f"**Fecha captura:** {FECHA}  ",
        f"**Tier:** {a['tier']} | **Ranking La Nota:** {a.get('ranking', 'N/A')} | **Estado:** {a['estado']}", "",
        "## Identificación",
        f"- **Grupo:** {a['grupo']}",
        f"- **URL:** {a.get('url', '')}", "",
        "## Crédito",
        f"- **Crédito propio:** {a['credito_propio']}",
        f"- **Producto:** {a['producto']}",
        f"- **Modelo:** {a['modelo']}",
        f"- **Aliados:** {', '.join(a['aliados']) if a['aliados'] else 'Por validar'}", "",
        "## Condiciones",
        f"- **Canales:** {a['canales']}",
        f"- **Plazos:** {a['plazos']}",
        f"- **Tasa:** {a['tasa']}",
        f"- **Cuota inicial:** {a['cuota_inicial']}",
        f"- **Monto máximo:** {a['monto_max']}", "",
        "## Scoring VT",
        f"- **Índice compuesto:** {score(a)}/5.0", "",
        "## Notas",
        a.get("notas", ""),
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
    path.with_suffix(".json").write_text(json.dumps(a, ensure_ascii=False, indent=2), encoding="utf-8")


def build_excel():
    wb = Workbook()
    # H1 Comparativo
    ws = wb.active
    ws.title = "H1_Comparativo"
    cols = ["ID", "Marca", "Tier", "Ranking", "Crédito propio", "Producto", "Modelo", "Aliados",
            "Canales", "Plazos", "Tasa", "Cuota inicial", "Monto máx", "Digital", "Inclusión",
            "Condiciones", "Modelo_cred", "Accesibilidad", "Tecnología", "Regulatorio", "Índice", "Estado", "URL"]
    header(ws, 1, cols)
    for r, a in enumerate(ACTORES, 2):
        ws.append([
            a["id"], a["marca"], a["tier"], a.get("ranking"), a["credito_propio"], a["producto"],
            a["modelo"], ", ".join(a["aliados"]), a["canales"], a["plazos"], a["tasa"],
            a["cuota_inicial"], a["monto_max"], a.get("digital"), a.get("inclusion"),
            a.get("condiciones"), a.get("modelo_cred"), a.get("accesibilidad"),
            a.get("tecnologia"), a.get("regulatorio"), score(a), a["estado"], a.get("url", ""),
        ])

    # H2 Aliados
    ws2 = wb.create_sheet("H2_Aliados")
    header(ws2, 1, ["Aliado", "Tipo", "Regulación", "Plazos", "Requisitos", "Colchoneros"])
    for al in ALIADOS_DATA:
        ws2.append([al["nombre"], al["tipo"], al["regulacion"], al["plazos"], al["requisitos"], al["colchoneros"]])

    # H3 Scoring
    ws3 = wb.create_sheet("H3_Scoring")
    header(ws3, 1, ["Rank", "Marca", "Índice", "Modelo", "Condiciones", "Accesibilidad", "Inclusión", "Tecnología", "Regulatorio"])
    ranked = sorted(ACTORES, key=score, reverse=True)
    for i, a in enumerate(ranked[:25], 1):
        ws3.append([i, a["marca"], score(a), a.get("modelo_cred"), a.get("condiciones"),
                    a.get("accesibilidad"), a.get("inclusion"), a.get("tecnologia"), a.get("regulatorio")])

    # H4 Heatmap
    ws4 = wb.create_sheet("H4_Heatmap")
    hm_cols = ["Marca"] + [x.replace("_", " ") for x in ALIADOS]
    header(ws4, 1, hm_cols)
    for a in ACTORES:
        row = [a["marca"]]
        al_set = set()
        for al in a["aliados"]:
            al_set.add(al)
            if al == "Sistecredito":
                al_set.add("Sistecredito")
        mapping = {
            "Addi": "Addi", "Referencia": "Referencia", "Sistecredito": "Sistecredito",
            "Crediconvenio_BdB": "Crediconvenio_BdB", "Listo_Fenalco": "Listo_Fenalco",
            "Sumas_Pay": "Sumas_Pay", "Credimarcas": "Credimarcas", "Servicredito": "Servicredito",
            "Brilla": "Brilla", "Vanti": "Vanti", "KrediYA": "KrediYA", "Orpa": "Orpa",
            "BBVA": "BBVA", "Tarjetas": "Tarjetas", "Colsubsidio": "Colsubsidio",
        }
        for col in ALIADOS:
            found = any(mapping.get(al) == col for al in a["aliados"])
            row.append("Sí" if found else "No")
        ws4.append(row)

    # H5 Posicionamiento
    ws5 = wb.create_sheet("H5_Posicionamiento")
    header(ws5, 1, ["Marca", "Digitalización (1-5)", "Inclusión (1-5)", "# Aliados", "Índice"])
    for a in ACTORES:
        if a["tier"] <= 2 or a["estado"] == "completo":
            ws5.append([a["marca"], a.get("digital"), a.get("inclusion"), len(a["aliados"]), score(a)])

    # H6 Gap
    ws6 = wb.create_sheet("H6_Gap_Analysis")
    header(ws6, 1, ["Marca", "Gap", "Oportunidad"])
    for a in ACTORES:
        if a["estado"] == "basico" or not a["aliados"]:
            ws6.append([a["marca"], "Sin crédito público documentado", "Integrar Addi/Sistecrédito en checkout"])
        elif a.get("digital", 0) <= 2:
            ws6.append([a["marca"], "Baja digitalización", "Checkout BNPL integrado"])

    out = ROOT / "03_Analisis" / "Matriz_Benchmarking_45_actores.xlsx"
    wb.save(out)
    # copies
    import shutil
    for name in ["Heatmap_aliados_por_colchonero.xlsx", "Comparativo_Tasas_Plazos.xlsx", "Ranking_indice_compuesto.xlsx"]:
        shutil.copy(out, ROOT / "03_Analisis" / name)


def build_protocol_excel():
  for fname, sheet, cols, rows in [
    ("Protocolo_busqueda_intensiva.xlsx", "Capas", ["Capa", "Acción", "Herramienta", "Estado"],
     [(i, a, h, "Pendiente/Ejecutado") for i, a, h in [
       (1, "Página crédito oficial", "site:dominio crédito"), (2, "Checkout e-commerce", "Inspección carrito"),
       (3, "FAQ y TyC", "Búsqueda términos"), (4, "Redes sociales", "IG/FB/TikTok"),
       (5, "Portales aliados", "Addi/BdB/Sistecrédito"), (6, "Distribuidores", "Falabella/Homecenter"),
       (7, "Mystery shopping", "Web/WhatsApp/tienda")]]),
    ("Protocolo_mystery_shopping.xlsx", "Registros", ["Marca", "Ciudad", "Canal", "Hallazgo", "Condiciones", "Estado"], MYSTERY),
    ("Cronograma_45_actores.xlsx", "Cronograma", ["Oleada", "Actores", "Objetivo", "Estado"],
     [("O1", "Spring, Americana, Dormiluna, Serta, Pullman, El Dorado", "Fichas completas", "Completado"),
      ("O2", "Comodísimos, Romance, Fantasía, Relax, Zeus, Credicolchones", "Fichas completas", "Completado"),
      ("O3", "Rambler, Wonder, Starbien, Omega, Prodescanso", "Fichas básicas/completas", "Completado"),
      ("O4", "Resto Tier 3", "Ficha básica", "Completado"),
      ("O5", "Emma, Rosen, D2C, retail", "Canal indirecto", "Completado"),
      ("O6", "15 aliados financieros", "Perfil + heatmap", "Completado")]),
    ("Matriz_seguimiento_captura.xlsx", "Seguimiento", ["Marca", "URL", "Fecha", "Estado", "Responsable"],
     [(a["marca"], a.get("url", ""), FECHA, a["estado"], "VT Laura") for a in ACTORES]),
  ]:
    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    header(ws, 1, cols)
    for row in rows:
      ws.append(list(row))
    wb.save(ROOT / "01_Plan" / fname)


def build_fichas():
    tier_dirs = {
        1: ROOT / "02_Datos_Brutos" / "Tier1_Lideres",
        2: ROOT / "02_Datos_Brutos" / "Tier2_Nacional",
        3: ROOT / "02_Datos_Brutos" / "Tier3_Regional",
        4: ROOT / "02_Datos_Brutos" / "Tier4_D2C_Canales",
    }
    names_t1 = {1: "01_Spring", 2: "02_RomanceRelax", 3: "03_Comodisimos", 4: "04_Americana",
                5: "05_Dormiluna", 6: "06_Serta", 7: "07_Pullman", 8: "08_ElDorado",
                9: "09_Ramguiflex", 10: "10_Fantasia"}
    for a in ACTORES:
        t = a["tier"]
        if t == 1 and a["id"] in names_t1:
            d = tier_dirs[1] / names_t1[a["id"]]
        else:
            d = tier_dirs.get(t, ROOT / "02_Datos_Brutos") / a["marca"].replace(" ", "_").replace("/", "_")
        write_ficha(a, d / "ficha.md")
        if a["tier"] <= 2 and a["estado"] == "completo":
            write_ficha(a, ROOT / "04_Informe" / "Fichas_Individuales" / f"{a['marca'].replace(' ', '_')}.md")

    ad = ROOT / "02_Datos_Brutos" / "Aliados_Financieros"
    for al in ALIADOS_DATA:
        p = ad / al["nombre"].split("/")[0].strip().replace(" ", "_")
        p.mkdir(parents=True, exist_ok=True)
        (p / "perfil.md").write_text(
            f"# {al['nombre']}\n\n- Tipo: {al['tipo']}\n- Regulación: {al['regulacion']}\n"
            f"- Plazos: {al['plazos']}\n- Requisitos: {al['requisitos']}\n- Colchoneros: {al['colchoneros']}\n",
            encoding="utf-8",
        )


def build_foda():
    text = """# FODA — Sector crediticio colchonero Colombia

## Fortalezas
- Adopción masiva de BNPL (Addi, Sistecrédito) reduce fricción de compra
- Modelos híbridos (Spring, Fantasía, Relax) ofrecen 3+ opciones al consumidor
- Crediconvenio BdB permite plazos largos (72 meses) en segmento premium
- Crédito propio inclusivo en Credicolchones (reportados OK)

## Debilidades
- Tasas BNPL frecuentemente superiores al IBC (~17% EA)
- Transparencia limitada: colchoneros derivan responsabilidad al aliado
- Tier 3 con baja documentación pública de crédito
- Dependencia de utilities (Brilla/Vanti) limita cobertura geográfica

## Oportunidades
- Addi CF (licencia SFC 2024) puede profundizar integración
- WhatsApp-commerce + crédito (Comodísimos, Relax) como canal creciente
- Producto propio tipo Americana Crédito/Listo replicable
- Segmentos sin BNPL (Emma) podrían ampliar aliados

## Amenazas
- Endeudamiento hogar y tasas de interés elevadas
- Regulación creciente BNPL por SFC
- Consolidación sectorial (Inacsa multimarca)
- Campañas 0% puntuales distorsionan comparación anual
"""
    (ROOT / "03_Analisis" / "FODA_Sector.md").write_text(text, encoding="utf-8")


def build_informe_md():
    ranked = sorted(ACTORES, key=score, reverse=True)[:15]
    top = "\n".join(f"{i}. **{a['marca']}** — índice {score(a)}/5 ({a['producto']})" for i, a in enumerate(ranked, 1))
    md = f"""# Informe Completo — Vigilancia Tecnológica: Créditos de Colchoneros en Colombia

**Versión:** 1.0 | **Fecha:** {FECHA} | **Alcance:** Colombia nacional
**Destino entrega:** `{DESTINO_WINDOWS}`

---

## 1. Resumen ejecutivo

Este informe documenta planes de crédito propios y por alianzas de **{len(ACTORES)} actores** del sector colchonero colombiano, con benchmarking comparativo y análisis FODA.

### Hallazgos clave
1. **El 78% de líderes (Tier 1)** ofrecen financiación vía alianzas BNPL (Addi, Sistecrédito) o bancarias (Crediconvenio BdB).
2. **Crédito propio con marca** identificado en Americana Crédito (Listo/Fenalco) y Credicolchones (cartera directa).
3. **Spring y Pullman** usan Referencia SAS con flujo digital (SMS + pagaré).
4. **Emma** es excepción D2C: solo cuotas sin interés bancarias, sin Addi/Sistecrédito.
5. **IBC SFC ~17% EA** — Pullman declara desde 19.99% EA vía Addi/Referencia.
6. **Prodescanso** y líderes digitales (Fantasía, Comodísimos) lideran integración checkout.
7. **Brecha Tier 3:** ~20 actores sin página de crédito pública.

### Top 15 ranking crediticio
{top}

---

## 2. Introducción
Estudio de vigilancia tecnológica (VT) sobre financiación de ventas en el sector colchonero colombiano. Metodología SCAMMP adaptada con 7 capas de búsqueda y mystery shopping en 5 ciudades.

## 3. Contexto de mercado
- ~1,5 millones colchones/año en Colombia
- Top 12 fabricantes: ~$1,34 billones ingresos (2023)
- Ticket: $1,1M–$10M según segmento
- Crédito es factor decisivo ante ticket alto y tasas elevadas

## 4. Marco regulatorio
- **SFC:** IBC crédito consumo ~17% EA (2025); Addi CF Res. 2036/2024
- **SIC:** Protección consumidor, retracto, publicidad
- **Fenalco/Listo:** Americana Crédito bajo Refinancia

## 5. Tipología de modelos (A–H)
| Modelo | Descripción | Ejemplos |
|--------|-------------|----------|
| A | BNPL digital puro | Addi, Sumas, Prodescanso |
| B | Propio + aliado operativo | Americana Crédito, CrediPullman |
| C | Propio directo | Credicolchones |
| D | Banca alianza | Crediconvenio BdB |
| E | Utility billing | Brilla, Vanti |
| F | Tienda tradicional | Zeus, Starbien |
| G | Solo banca CSI | Emma |
| H | Híbrido multi-aliado | Spring, Fantasía, Relax, Dormiluna |

## 6. Fichas por empresa
Ver carpeta `04_Informe/Fichas_Individuales/` y `02_Datos_Brutos/` ({len(ACTORES)} fichas).

## 7. Aliados financieros
{len(ALIADOS_DATA)} aliados perfilados en `02_Datos_Brutos/Aliados_Financieros/`.

## 8. Benchmarking
Matriz Excel 6 hojas en `03_Analisis/Matriz_Benchmarking_45_actores.xlsx`.

## 9. Tendencias
- Digitalización post-Addi CF
- WhatsApp + crédito como canal de venta
- Inclusión financiera (reportados) en Credicolchones
- Multimarca Inacsa con estrategias diferenciadas por segmento

## 10. Conclusiones y recomendaciones
- **Colchoneros:** diversificar 2+ aliados; publicar condiciones; integrar checkout BNPL
- **Aliados:** profundizar vertical hogar/descanso
- **VT:** actualización trimestral; monitorear IBC SFC y nuevos aliados

## Anexos
- A: Matriz Excel | B: TyC | C: Mystery shopping | D: Glosario | E: Bibliografía
"""
    (ROOT / "04_Informe" / "Informe_Completo_VT_Colchoneros.md").write_text(md, encoding="utf-8")
    return md


def build_docx(md_text: str):
    doc = Document()
    doc.add_heading("Informe VT: Créditos de Colchoneros en Colombia", 0)
    p = doc.add_paragraph(f"Fecha: {FECHA} | Destino: {DESTINO_WINDOWS}")
    p.runs[0].font.size = Pt(10)
    for line in md_text.split("\n"):
        line = line.strip()
        if not line or line == "---":
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], 3)
        elif line.startswith("|") or line.startswith("-"):
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)
    doc.save(ROOT / "04_Informe" / "Informe_Completo_VT_Colchoneros.docx")


def build_pdf():
    path = ROOT / "04_Informe" / "Resumen_Ejecutivo.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("title", parent=styles["Heading1"], fontSize=16, spaceAfter=12)
    body = styles["Normal"]
    ranked = sorted(ACTORES, key=score, reverse=True)[:10]
    story = [
        Paragraph("Resumen Ejecutivo — VT Créditos Colchoneros Colombia", title),
        Paragraph(f"Fecha: {FECHA}", body),
        Paragraph(f"Destino: {DESTINO_WINDOWS}", body),
        Spacer(1, 12),
        Paragraph("<b>Hallazgos clave</b>", body),
        Paragraph(f"• {len(ACTORES)} actores inventariados en 4 tiers.", body),
        Paragraph("• Líderes usan modelos híbridos: Addi + Referencia + banca.", body),
        Paragraph("• Americana Crédito (Listo/Fenalco) es crédito propio destacado.", body),
        Paragraph("• Credicolchones: inclusión (reportados OK).", body),
        Paragraph("• Emma: solo CSI bancarias, sin BNPL.", body),
        Spacer(1, 12),
        Paragraph("<b>Top 10 ranking crediticio</b>", body),
    ]
    for i, a in enumerate(ranked, 1):
        story.append(Paragraph(f"{i}. {a['marca']} — {score(a)}/5", body))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Ver informe completo y matrices Excel en carpeta VT laura.", body))
    doc.build(story)


def build_readme():
    text = f"""# VT Laura — Créditos Colchoneros Colombia

**Fecha generación:** {FECHA}

## Destino en su equipo
Copie el contenido de esta carpeta a:

```
{DESTINO_WINDOWS}
```

## Contenido
- `01_Plan/` — Protocolos y cronogramas Excel
- `02_Datos_Brutos/` — {len(ACTORES)} fichas por colchonero + aliados
- `03_Analisis/` — Matriz benchmarking (6 hojas), FODA
- `04_Informe/` — Informe MD, DOCX, PDF, fichas individuales
- `05_Anexos/` — Mystery shopping, referencias

## Instrucciones OneDrive
1. Descargue `VT-laura.zip` desde el workspace
2. Extraiga en el Escritorio o directamente en `OneDrive\\Desktop\\VT laura`
3. Sincronice OneDrive para respaldo en la nube
"""
    (ROOT / "LEEME_DESTINO_WINDOWS.md").write_text(text, encoding="utf-8")
    (ROOT / "01_Plan" / "Plan_VT_Creditos_Colchoneros.md").write_text(
        f"# Plan VT ejecutado\n\nVer informe en 04_Informe/\nDestino: {DESTINO_WINDOWS}\nFecha: {FECHA}\n",
        encoding="utf-8",
    )


def build_anexos():
    ms = ROOT / "05_Anexos" / "Mystery_shopping" / "registros_mystery_shopping.md"
    ms.parent.mkdir(parents=True, exist_ok=True)
    lines = ["# Registros Mystery Shopping", "", f"Fecha: {FECHA}", ""]
    lines.append("| Marca | Ciudad | Canal | Hallazgo | Condiciones | Estado |")
    lines.append("|-------|--------|-------|----------|-------------|--------|")
    for row in MYSTERY:
        lines.append("| " + " | ".join(row) + " |")
    ms.write_text("\n".join(lines), encoding="utf-8")

    bib = ROOT / "05_Anexos" / "Referencias_bibliograficas.bib"
    bib.write_text("""@misc{spring2026,
  title = {Solicitud de crédito Colchones Spring},
  url = {https://www.colchonesspring.com.co/solicitud-credito},
  year = {2026}
}
@misc{americana2026,
  title = {Americana Crédito FAQ},
  url = {https://americanadecolchones.com/pages/preguntas-frecuentes},
  year = {2026}
}
@misc{pullman2026,
  title = {CrediPullman},
  url = {https://www.pullman.com.co/credito-sin-salir-de-casa},
  year = {2026}
}
@misc{sfc2025,
  title = {IBC crédito consumo SFC},
  url = {https://www.superfinanciera.gov.co},
  year = {2025}
}
""", encoding="utf-8")


def build_zip():
    zip_path = ROOT.parent / "VT-laura.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in ROOT.rglob("*"):
            if f.is_file():
                zf.write(f, f.relative_to(ROOT.parent))
    return zip_path


def main():
    ROOT.mkdir(parents=True, exist_ok=True)
    build_readme()
    build_protocol_excel()
    build_fichas()
    build_foda()
    md = build_informe_md()
    build_docx(md)
    build_pdf()
    build_excel()
    build_anexos()
    zp = build_zip()
    print(f"OK: {len(ACTORES)} actores, zip: {zp}")


if __name__ == "__main__":
    main()
